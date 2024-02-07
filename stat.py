# Importowanie bibliotek
import pandas as pd
import numpy as np
import scipy.stats as stats
import matplotlib.pyplot as plt

# Wczytywanie danych
data = pd.read_excel("data.xlsx")
plec = data["Płeć"]
rodzic = data["Rodzic"]
swls = data["SWLS"]
kss = data["KSS"]

# Obliczanie liczby obserwacji
n = len(data)

# Korelacja Spearmana
spearman = stats.spearmanr(swls, kss)
print(f"Korelacja rang Spearmana między SWLS a KSS wynosi {spearman[0]:.2f}, p-wartość wynosi {spearman[1]:.4f}.")

# Sprawdzamy, czy rozkład jest normalny w obu grupach
if n < 50:
    # Dla małych próbek stosujemy test Shapiro-Wilka
    shapiro_swls_rodzic = stats.shapiro(swls[rodzic == 1])
    shapiro_swls_nierodzic = stats.shapiro(swls[rodzic == 0])
    shapiro_kss_rodzic = stats.shapiro(kss[rodzic == 1])
    shapiro_kss_nierodzic = stats.shapiro(kss[rodzic == 0])
    print(f"Wyniki testu Shapiro-Wilka dla SWLS w grupie rodziców: statystyka = {shapiro_swls_rodzic[0]:.2f}, p-wartość = {shapiro_swls_rodzic[1]:.4f}.")
    print(f"Wyniki testu Shapiro-Wilka dla SWLS w grupie nierodziców: statystyka = {shapiro_swls_nierodzic[0]:.2f}, p-wartość = {shapiro_swls_nierodzic[1]:.4f}.")
    print(f"Wyniki testu Shapiro-Wilka dla KSS w grupie rodziców: statystyka = {shapiro_kss_rodzic[0]:.2f}, p-wartość = {shapiro_kss_rodzic[1]:.4f}.")
    print(f"Wyniki testu Shapiro-Wilka dla KSS w grupie nierodziców: statystyka = {shapiro_kss_nierodzic[0]:.2f}, p-wartość = {shapiro_kss_nierodzic[1]:.4f}.")
else:
    # Dla dużych próbek stosujemy test Jarque-Bera
    jarque_swls_rodzic = stats.jarque_bera(swls[rodzic == 1])
    jarque_swls_nierodzic = stats.jarque_bera(swls[rodzic == 0])
    jarque_kss_rodzic = stats.jarque_bera(kss[rodzic == 1])
    jarque_kss_nierodzic = stats.jarque_bera(kss[rodzic == 0])
    print(f"Wyniki testu Jarque-Bera dla SWLS w grupie rodziców: statystyka = {jarque_swls_rodzic[0]:.2f}, p-wartość = {jarque_swls_rodzic[1]:.4f}.")
    print(f"Wyniki testu Jarque-Bera dla SWLS w grupie nierodziców: statystyka = {jarque_swls_nierodzic[0]:.2f}, p-wartość = {jarque_swls_nierodzic[1]:.4f}.")
    print(f"Wyniki testu Jarque-Bera dla KSS w grupie rodziców: statystyka = {jarque_kss_rodzic[0]:.2f}, p-wartość = {jarque_kss_rodzic[1]:.4f}.")
    print(f"Wyniki testu Jarque-Bera dla KSS w grupie nierodziców: statystyka = {jarque_kss_nierodzic[0]:.2f}, p-wartość = {jarque_kss_nierodzic[1]:.4f}.")

# Wyjaśniamy, jak interpretować p-wartość testu normalności
print("Jeśli p-wartość testu normalności jest mniejsza niż 0.05, to odrzucamy hipotezę, że rozkład jest normalny. Jeśli p-wartość jest większa lub równa 0.05, to nie mamy podstaw do odrzucenia hipotezy o normalności rozkładu.")

# Sprawdzamy, czy wariancje są równe w obu grupach
levene_swls = stats.levene(swls[rodzic == 1], swls[rodzic == 0])
levene_kss = stats.levene(kss[rodzic == 1], kss[rodzic == 0])
print(f"Wyniki testu Levene'a dla SWLS w grupach rodziców i nierodziców: statystyka = {levene_swls[0]:.2f}, p-wartość = {levene_swls[1]:.4f}.")
print(f"Wyniki testu Levene'a dla KSS w grupach rodziców i nierodziców: statystyka = {levene_kss[0]:.2f}, p-wartość = {levene_kss[1]:.4f}.")

# Wyjaśniamy, jak interpretować p-wartość testu równości wariancji
print("Jeśli p-wartość testu równości wariancji jest mniejsza niż 0.05, to odrzucamy hipotezę, że wariancje są równe. Jeśli p-wartość jest większa lub równa 0.05, to nie mamy podstaw do odrzucenia hipotezy o równości wariancji.")

# Wybór testu t-Studenta lub Mann-Whitney'a U z uwzględnieniem normalności i równości wariancji
if (jarque_swls_rodzic[1] > 0.05 and jarque_swls_nierodzic[1] > 0.05 and jarque_kss_rodzic[1] > 0.05 and jarque_kss_nierodzic[1] > 0.05):
    # Jeśli obie zmienne są normalnie rozłożone w obu grupach, sprawdzamy równość wariancji
    if (levene_swls[1] >= 0.05 and levene_kss[1] >= 0.05):
        # Jeśli wariancje są równe, stosujemy test t-Studenta bez korekty
        ttest_swls = stats.ttest_ind(swls[rodzic == 1], swls[rodzic == 0], equal_var=True)
        ttest_kss = stats.ttest_ind(kss[rodzic == 1], kss[rodzic == 0], equal_var=True)
        print(f"Wyniki testu t-Studenta dla SWLS w grupach rodziców i nierodziców: statystyka = {ttest_swls[0]:.2f}, p-wartość = {ttest_swls[1]:.4f}.")
        print(f"Wyniki testu t-Studenta dla KSS w grupach rodziców i nierodziców: statystyka = {ttest_kss[0]:.2f}, p-wartość = {ttest_kss[1]:.4f}.")
        # Wyjaśniamy, jak interpretować p-wartość testu t-Studenta
        print("Jeśli p-wartość testu t-Studenta jest mniejsza niż 0.05, to odrzucamy hipotezę, że średnie są równe w obu grupach. Jeśli p-wartość jest większa lub równa 0.05, to nie mamy podstaw do odrzucenia hipotezy o równości średnich.")
    else:
        # Jeśli wariancje są różne, stosujemy test t-Studenta z korektą Welcha
        ttest_swls = stats.ttest_ind(swls[rodzic == 1], swls[rodzic == 0], equal_var=False)
        ttest_kss = stats.ttest_ind(kss[rodzic == 1], kss[rodzic == 0], equal_var=False)
        print(f"Wyniki testu t-Studenta z korektą Welcha dla SWLS w grupach rodziców i nierodziców: statystyka = {ttest_swls[0]:.2f}, p-wartość = {ttest_swls[1]:.4f}.")
        print(f"Wyniki testu t-Studenta z korektą Welcha dla KSS w grupach rodziców i nierodziców: statystyka = {ttest_kss[0]:.2f}, p-wartość = {ttest_kss[1]:.4f}.")
        # Wyjaśniamy, jak interpretować p-wartość testu t-Studenta z korektą Welcha
        print("Jeśli p-wartość testu t-Studenta z korektą Welcha jest mniejsza niż 0.05, to odrzucamy hipotezę, że średnie są równe w obu grupach. Jeśli p-wartość jest większa lub równa 0.05, to nie mamy podstaw do odrzucenia hipotezy o równości średnich.")
else:
    # W przeciwnym razie, sprawdzamy równość wariancji
    if (levene_swls[1] >= 0.05 and levene_kss[1] >= 0.05):
        # Jeśli wariancje są równe, stosujemy test Mann-Whitney'a U bez korekty
        mannwhitney_swls = stats.mannwhitneyu(swls[rodzic == 1], swls[rodzic == 0], alternative="two-sided")
        mannwhitney_kss = stats.mannwhitneyu(kss[rodzic == 1], kss[rodzic == 0], alternative="two-sided")
        print(f"Wyniki testu Mann-Whitney'a U dla SWLS w grupach rodziców i nierodziców: statystyka = {mannwhitney_swls[0]:.2f}, p-wartość = {mannwhitney_swls[1]:.4f}.")
        print(f"Wyniki testu Mann-Whitney'a U dla KSS w grupach rodziców i nierodziców: statystyka = {mannwhitney_kss[0]:.2f}, p-wartość = {mannwhitney_kss[1]:.4f}.")
        # Wyjaśniamy, jak interpretować p-wartość testu Mann-Whitney'a U
        print("Jeśli p-wartość testu Mann-Whitney'a U jest mniejsza niż 0.05, to odrzucamy hipotezę, że mediany są równe w obu grupach. Jeśli p-wartość jest większa lub równa 0.05, to nie mamy podstaw do odrzucenia hipotezy o równości median.")
    else:
        # Jeśli wariancje są różne, stosujemy test Mann-Whitney'a U z korektą Ties
        mannwhitney_swls = stats.mannwhitneyu(swls[rodzic == 1], swls[rodzic == 0], alternative="two-sided", use_continuity=False)
        mannwhitney_kss = stats.mannwhitneyu(kss[rodzic == 1], kss[rodzic == 0], alternative="two-sided", use_continuity=False)
        print(f"Wyniki testu Mann-Whitney'a U z korektą Ties dla SWLS w grupach rodziców i nierodziców: statystyka = {mannwhitney_swls[0]:.2f}, p-wartość = {mannwhitney_swls[1]:.4f}.")
        print(f"Wyniki testu Mann-Whitney'a U z korektą Ties dla KSS w grupach rodziców i nierodziców: statystyka = {mannwhitney_kss[0]:.2f}, p-wartość = {mannwhitney_kss[1]:.4f}.")
        # Wyjaśniamy, jak interpretować p-wartość testu Mann-Whitney'a U z korektą Ties
        print("Jeśli p-wartość testu Mann-Whitney'a U z korektą Ties jest mniejsza niż 0.05, to odrzucamy hipotezę, że mediany są równe w obu grupach. Jeśli p-wartość jest większa lub równa 0.05, to nie mamy podstaw do odrzucenia hipotezy o równości median.")




# Wykresy Q-Q
fig, ax = plt.subplots(2, 2, figsize=(10, 10))
stats.probplot(swls[rodzic == 1], dist="norm", plot=ax[0, 0])
ax[0, 0].set_title("Wykres Q-Q dla SWLS w grupie rodziców")
stats.probplot(swls[rodzic == 0], dist="norm", plot=ax[0, 1])
ax[0, 1].set_title("Wykres Q-Q dla SWLS w grupie nierodziców")
stats.probplot(kss[rodzic == 1], dist="norm", plot=ax[1, 0])
ax[1, 0].set_title("Wykres Q-Q dla KSS w grupie rodziców")
stats.probplot(kss[rodzic == 0], dist="norm", plot=ax[1, 1])
ax[1, 1].set_title("Wykres Q-Q dla KSS w grupie nierodziców")
plt.tight_layout()
plt.show()

# Zapisywanie wyników
# Tworzenie DataFrame z wynikami analizy
wyniki = pd.DataFrame({"Zmienna": ["SWLS", "KSS"], 
                       "Korelacja Spearmana": [spearman[0], spearman[0]], 
                       "p-wartość korelacji": [spearman[1], spearman[1]], 
                       "Test normalności rodziców": ["Shapiro-Wilka" if n < 50 else "Jarque-Bera", "Shapiro-Wilka" if n < 50 else "Jarque-Bera"], 
                       "p-wartość normalności rodziców": [jarque_swls_rodzic[1] if n >= 50 else shapiro_swls_rodzic[1], jarque_kss_rodzic[1] if n >= 50 else shapiro_kss_rodzic[1]], 
                       "Test normalności nierodziców": ["Shapiro-Wilka" if n < 50 else "Jarque-Bera", "Shapiro-Wilka" if n < 50 else "Jarque-Bera"], 
                       "p-wartość normalności nierodziców": [jarque_swls_nierodzic[1] if n >= 50 else shapiro_swls_nierodzic[1], jarque_kss_nierodzic[1] if n >= 50 else shapiro_kss_nierodzic[1]], 
                       "Test porównania grup": ["t-Studenta" if (jarque_swls_rodzic[1] > 0.05 and jarque_swls_nierodzic[1] > 0.05 and jarque_kss_rodzic[1] > 0.05 and jarque_kss_nierodzic[1] > 0.05) else "Mann-Whitney'a U", "t-Studenta" if (jarque_swls_rodzic[1] > 0.05 and jarque_swls_nierodzic[1] > 0.05 and jarque_kss_rodzic[1] > 0.05 and jarque_kss_nierodzic[1] > 0.05) else "Mann-Whitney'a U"], 
                       "p-wartość porównania grup": [ttest_swls[1] if (jarque_swls_rodzic[1] > 0.05 and jarque_swls_nierodzic[1] > 0.05 and jarque_kss_rodzic[1] > 0.05 and jarque_kss_nierodzic[1] > 0.05) else mannwhitney_swls[1], ttest_kss[1] if (jarque_swls_rodzic[1] > 0.05 and jarque_swls_nierodzic[1] > 0.05 and jarque_kss_rodzic[1] > 0.05 and jarque_kss_nierodzic[1] > 0.05) else mannwhitney_kss[1]]})
# Zapisywanie wyników do pliku wyniki.xlsx
wyniki.to_excel("wyniki.xlsx", index=False)
# Tworzenie tabeli w pliku wyniki.docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
doc = Document()
doc.add_heading("Wyniki analizy danych", level=1)
table = doc.add_table(rows=3, cols=8)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
# Dodawanie nagłówków tabeli
table.cell(0, 0).text = "Zmienna"
table.cell(0, 1).text = "Korelacja Spearmana"
table.cell(0, 2).text = "p-wartość korelacji"
table.cell(0, 3).text = "Test normalności rodziców"
table.cell(0, 4).text = "p-wartość normalności rodziców"
table.cell(0, 5).text = "Test normalności nierodziców"
table.cell(0, 6).text = "p-wartość normalności nierodziców"
table.cell(0, 7).text = "Test porównania grup"
table.cell(0, 8).text = "p-wartość porównania grup"
# Dodawanie danych tabeli
for i in range(2):
    for j in range(8):
        table.cell(i+1, j).text = str(wyniki.iloc[i, j])
# Ustawianie szerokości kolumn
for cell in table.rows[0].cells:
    cell.width = Cm(2.5)
# Wyśrodkowywanie tekstu w komórkach
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# Zapisywanie dokumentu
doc.save("wyniki.docx")
